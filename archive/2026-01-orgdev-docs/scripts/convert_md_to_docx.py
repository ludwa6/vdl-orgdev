import re
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path, img_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return

    doc = Document()
    
    # Set narrow margins for better image/table fit
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_table = False
    table_data = []
    in_mermaid = False

    for line in lines:
        stripped = line.strip()

        # Skip Mermaid Code Blocks and insert image instead
        if stripped.startswith('```mermaid'):
            in_mermaid = True
            # Insert the diagram image instead of the code
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(7)) # Fit to page
                doc.add_paragraph("Figure: Application Schema Entity-Relationship Diagram").alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        
        if in_mermaid:
            if stripped == '```':
                in_mermaid = False
            continue

        # Skip standard code blocks for now or handle them simply
        if stripped.startswith('```') and not in_mermaid:
            continue

        # Handle Tables
        if '|' in line:
            if not in_table:
                in_table = True
                table_data = []
            
            if re.match(r'^[|:\s-]+$', stripped):
                continue
            
            cells = [c.strip() for c in stripped.split('|') if c.strip() or stripped.startswith('|')]
            if stripped.startswith('|'): cells = [c.strip() for c in stripped.split('|')][1:-1]
            else: cells = [c.strip() for c in stripped.split('|')]
            
            if any(cells):
                table_data.append(cells)
            continue
        else:
            if in_table:
                if table_data:
                    cols = max(len(row) for row in table_data)
                    table = doc.add_table(rows=len(table_data), cols=cols)
                    table.style = 'Table Grid'
                    for i, row in enumerate(table_data):
                        for j, cell_text in enumerate(row):
                            if j < cols:
                                table.cell(i, j).text = cell_text
                in_table = False
                table_data = []

        # Handle Headers
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        
        # Handle Bullet Points
        elif stripped.startswith('* ') or stripped.startswith('- '):
            doc.add_paragraph(stripped[2:], style='List Bullet')
        
        # Handle simple bolding
        elif stripped:
            p = doc.add_paragraph()
            # Basic bolding for **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph()

    doc.save(docx_path)
    print(f"Successfully saved to {docx_path}")

if __name__ == "__main__":
    BASE_DIR = "/Users/walt/.gemini/antigravity/brain/9d3c833e-2993-45ea-b6bf-45a93fc4e691"
    MD_FILE = os.path.join(BASE_DIR, "schema_documentation.md")
    DOCX_FILE = os.path.join(BASE_DIR, "schema_documentation.docx")
    IMG_FILE = os.path.join(BASE_DIR, "erd_diagram_1770460022393.png")
    convert_md_to_docx(MD_FILE, DOCX_FILE, IMG_FILE)
