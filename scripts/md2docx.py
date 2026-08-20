#!/usr/bin/env python3
"""markdown -> docx, for the handbook outputs.

The direction of this pipeline was reversed on 2026-08-20. Markdown in this
repo is now the source of the Handbook; .docx, .pages and .pdf are outputs
generated from it. Before that date the .pages files were the source and
scripts/docx2md.py rendered them into this repo as a read-only audit trail.

That inversion is the reason this script exists. It is the mirror of
docx2md.py, and the two are deliberately symmetrical:

    docx2md.py   24pt -> #     18pt -> ##     14pt -> ###
    md2docx.py   #    -> 24pt  ##   -> 18pt   ###  -> 14pt

so that a document round-tripped through both comes back to the same
markdown. That property is worth keeping: it is what lets us accept an
edit that arrives as a Word document without hand-transcribing it.

Usage:
    scripts/md2docx.py handbook/pt/04_Como_Trabalhamos.md
    scripts/md2docx.py handbook/pt/*.md handbook/en/*.md
    scripts/md2docx.py --all

The .docx is written beside its .md. To produce .pages or .pdf, open the
.docx in Pages and export — Pages has no scriptable "export as Pages" verb,
so that leg stays manual. See handbook/README.md.

Deliberately not supported: tables of contents, styled cover pages, or
Nita's Pages design. This produces a clean, readable, correctly-structured
document, not a designed one. If the organisation wants the designed look
back, the answer is a .docx template with named styles, not a richer script.
"""
import sys
import glob
import re
from pathlib import Path

import docx
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Mirror of docx2md.py's thresholds, in EMU, so the round-trip is exact.
H1, H2, H3 = 304800, 228600, 177800
BODY = Pt(11)

REPO = Path(__file__).resolve().parent.parent

# `**bold**` and `*italic*`, non-greedy, longest marker first.
INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def add_runs(paragraph, text):
    """Write text into a paragraph, honouring ** and * inline markers."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            run = paragraph.add_run(piece[1:-1])
            run.italic = True
        else:
            paragraph.add_run(piece)


def heading(doc, text, emu):
    p = doc.add_paragraph()
    add_runs(p, text)
    for run in p.runs:
        run.font.size = Emu(emu)
        run.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    add_runs(p, text)
    for run in p.runs:
        if run.font.size is None:
            run.font.size = BODY
    p.paragraph_format.space_after = Pt(8)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    add_runs(p, text)
    for run in p.runs:
        run.font.size = BODY
    p.paragraph_format.space_after = Pt(4)
    return p


def table(doc, rows):
    """rows: list of lists of cell strings; first row is the header."""
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for cell, text in zip(cells, row):
            cell.paragraphs[0].text = ""
            add_runs(cell.paragraphs[0], text)
            for run in cell.paragraphs[0].runs:
                run.font.size = BODY
                if i == 0:
                    run.bold = True
    doc.add_paragraph()
    return t


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator(line):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def convert(md_path: Path) -> Path:
    lines = md_path.read_text(encoding="utf-8").split("\n")
    doc = docx.Document()

    # Give the default style a sane body size; headings override per-run.
    doc.styles["Normal"].font.size = BODY

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Table: a pipe row followed by a separator row.
        if stripped.startswith("|") and i + 1 < len(lines) and is_separator(lines[i + 1]):
            rows = [split_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            table(doc, rows)
            continue

        if stripped.startswith("### "):
            heading(doc, stripped[4:], H3)
        elif stripped.startswith("## "):
            heading(doc, stripped[3:], H2)
        elif stripped.startswith("# "):
            heading(doc, stripped[2:], H1)
        elif stripped.startswith("- "):
            bullet(doc, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\. ", "", stripped))
            for run in p.runs:
                run.font.size = BODY
        else:
            body(doc, stripped)
        i += 1

    out = md_path.with_suffix(".docx")
    doc.save(out)
    return out


def main(argv):
    if len(argv) == 1 or argv[1] in ("-h", "--help"):
        print(__doc__)
        return 0

    if argv[1] == "--all":
        paths = sorted(REPO.glob("handbook/pt/*.md")) + sorted(REPO.glob("handbook/en/*.md"))
    else:
        paths = []
        for a in argv[1:]:
            paths.extend(Path(p) for p in glob.glob(a))

    # README.md and the translation working notes are not Handbook documents.
    paths = [p for p in paths if p.name != "README.md" and not p.name.startswith("00_")]

    if not paths:
        print("nothing to convert", file=sys.stderr)
        return 1

    for p in paths:
        out = convert(p)
        print(f"  {p.relative_to(REPO)} -> {out.relative_to(REPO)}")
    print(f"\n{len(paths)} document(s) generated.")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
